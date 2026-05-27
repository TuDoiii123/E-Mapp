"""
Template filler — điền thông tin vào file Word template dùng python-docx.
Placeholder format: {{field_name}}

Dùng chiến lược "gộp toàn bộ text trong paragraph rồi thay thế" để tránh
bị split placeholder trên nhiều Run.
"""
import os
import re
import logging
from io import BytesIO

log = logging.getLogger('template_filler')

_PLACEHOLDER_RE = re.compile(r'\{\{(\w+)\}\}')


def _replace_in_paragraph(para, fields: dict[str, str]) -> None:
    """
    Thay thế tất cả {{key}} trong một paragraph.
    Giữ nguyên font/bold/italic của Run đầu tiên.
    """
    full_text = ''.join(r.text for r in para.runs)
    if '{{' not in full_text:
        return

    for key, value in fields.items():
        placeholder = f'{{{{{key}}}}}'
        full_text = full_text.replace(placeholder, str(value) if value is not None else '')

    # Đặt toàn bộ text vào Run đầu, xóa các Run còn lại
    if para.runs:
        para.runs[0].text = full_text
        for run in para.runs[1:]:
            run.text = ''


def fill_template(template_path: str, fields: dict) -> BytesIO:
    """
    Điền thông tin vào Word template.

    Args:
        template_path: Đường dẫn tới file .docx template
        fields: Dict {field_name: value}

    Returns:
        BytesIO chứa file .docx đã điền

    Raises:
        ImportError: nếu python-docx chưa cài
        FileNotFoundError: nếu template không tồn tại
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError('python-docx chưa được cài. Chạy: pip install python-docx')

    if not os.path.exists(template_path):
        raise FileNotFoundError(f'Template không tồn tại: {template_path}')

    # Ép tất cả values sang str
    str_fields = {k: (str(v) if v is not None else '') for k, v in fields.items()}

    doc = Document(template_path)

    # Paragraphs ngoài table
    for para in doc.paragraphs:
        _replace_in_paragraph(para, str_fields)

    # Paragraphs trong tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, str_fields)

    # Headers / Footers
    for section in doc.sections:
        for para in section.header.paragraphs:
            _replace_in_paragraph(para, str_fields)
        for para in section.footer.paragraphs:
            _replace_in_paragraph(para, str_fields)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def get_template_fields(template_path: str) -> list[str]:
    """
    Quét file Word và trả về tất cả tên field {{field_name}} tìm thấy.
    """
    try:
        from docx import Document
    except ImportError:
        return []

    if not os.path.exists(template_path):
        return []

    doc = Document(template_path)
    fields: set[str] = set()

    def _scan(para):
        full = ''.join(r.text for r in para.runs)
        for m in _PLACEHOLDER_RE.finditer(full):
            fields.add(m.group(1))

    for para in doc.paragraphs:
        _scan(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _scan(para)

    return sorted(fields)


def add_placeholders_to_template(template_path: str, output_path: str,
                                  field_map: dict[str, str]) -> None:
    """
    Tiện ích: thay thế các giá trị mẫu trong template gốc bằng {{placeholder}}.
    Dùng để chuẩn bị template từ file Word có sẵn.

    Args:
        template_path: File Word gốc
        output_path:   File Word output với placeholders
        field_map:     {text_trong_file_gốc: tên_field} ví dụ {"Nguyễn Văn A": "ho_ten"}
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError('python-docx chưa được cài')

    doc = Document(template_path)

    def _inject(para):
        full = ''.join(r.text for r in para.runs)
        for original, field_name in field_map.items():
            full = full.replace(original, f'{{{{{field_name}}}}}')
        if para.runs:
            para.runs[0].text = full
            for r in para.runs[1:]:
                r.text = ''

    for para in doc.paragraphs:
        _inject(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _inject(para)

    doc.save(output_path)
    log.info(f'Đã tạo template với placeholders: {output_path}')
