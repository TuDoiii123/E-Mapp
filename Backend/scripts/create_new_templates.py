"""
Tao them cac mau giay to moi - ca .docx va .pdf.
"""
import sys, os
sys.path.insert(0, '.')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

TDIR = os.path.normpath(os.path.join(os.path.dirname(__file__), '..', 'data', 'templates'))
FONT_REGULAR = r'C:\Windows\Fonts\arial.ttf'
FONT_BOLD    = r'C:\Windows\Fonts\arialbd.ttf'


# ── Helper docx ───────────────────────────────────────────────────────────────
def new_doc():
    doc = Document()
    s = doc.sections[0]
    s.left_margin = Cm(2.5); s.right_margin = Cm(2)
    s.top_margin = Cm(2);    s.bottom_margin = Cm(1.5)
    return doc


def h_center(doc, text, size=12, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)
    return p


def line(doc, label, width=40, bold_label=True):
    p = doc.add_paragraph()
    r = p.add_run(label); r.bold = bold_label; r.font.size = Pt(11)
    p.add_run(' ' + '_' * width).font.size = Pt(11)


def sig_block(doc, role='NGƯỜI LÀM ĐƠN', date_text='........, ngày ...... tháng ...... năm ......'):
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(date_text).italic = True
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(role).bold = True
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('(Ký, ghi rõ họ tên)').italic = True


def header_vn(doc):
    h_center(doc, 'CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM', 12, True)
    h_center(doc, 'Độc lập - Tự do - Hạnh phúc', 12, True)
    h_center(doc, '─────────────────────', 11)


def title(doc, text, sub=None):
    doc.add_paragraph()
    h_center(doc, text, 13, True)
    if sub:
        h_center(doc, sub, 10)
    doc.add_paragraph()


def save(doc, fname):
    path = os.path.join(TDIR, fname)
    doc.save(path)
    print(f'  [docx] {fname:55s} {os.path.getsize(path)//1024}KB')
    return path


# ── Helper PDF (fpdf2) ────────────────────────────────────────────────────────
from fpdf import FPDF

class VnPDF(FPDF):
    def __init__(self):
        super().__init__()
        self.add_font('Arial', '', FONT_REGULAR)
        self.add_font('Arial', 'B', FONT_BOLD)
        self.set_auto_page_break(auto=True, margin=15)
        self.set_margins(25, 20, 20)

    def vheader(self, title_text, sub_text=None):
        self.set_font('Arial', 'B', 12)
        self.cell(0, 7, 'CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM', align='C', new_x='LMARGIN', new_y='NEXT')
        self.cell(0, 7, 'Độc lập - Tự do - Hạnh phúc', align='C', new_x='LMARGIN', new_y='NEXT')
        self.set_font('Arial', '', 11)
        self.cell(0, 5, '──────────────────────────────', align='C', new_x='LMARGIN', new_y='NEXT')
        self.ln(3)
        self.set_font('Arial', 'B', 13)
        self.cell(0, 8, title_text, align='C', new_x='LMARGIN', new_y='NEXT')
        if sub_text:
            self.set_font('Arial', '', 10)
            self.cell(0, 6, sub_text, align='C', new_x='LMARGIN', new_y='NEXT')
        self.ln(4)

    def field(self, label, width=120):
        self.set_font('Arial', 'B', 11)
        self.cell(55, 7, label)
        self.set_font('Arial', '', 11)
        self.cell(width, 7, '_' * int(width / 2.2), new_x='LMARGIN', new_y='NEXT')

    def para(self, text, bold=False):
        self.set_font('Arial', 'B' if bold else '', 11)
        self.multi_cell(0, 6, text, new_x='LMARGIN', new_y='NEXT')

    def sig(self, role='NGƯỜI LÀM ĐƠN'):
        self.ln(6)
        self.set_font('Arial', '', 10)
        self.cell(0, 6, '........, ngày ...... tháng ...... năm ......', align='R', new_x='LMARGIN', new_y='NEXT')
        self.set_font('Arial', 'B', 11)
        self.cell(0, 6, role, align='R', new_x='LMARGIN', new_y='NEXT')
        self.set_font('Arial', '', 10)
        self.cell(0, 6, '(Ký, ghi rõ họ tên)', align='R', new_x='LMARGIN', new_y='NEXT')


def save_pdf(pdf, fname):
    path = os.path.join(TDIR, fname)
    pdf.output(path)
    print(f'  [pdf]  {fname:55s} {os.path.getsize(path)//1024}KB')


# ════════════════════════════════════════════════════════════════════════════════
# DOCX TEMPLATES
# ════════════════════════════════════════════════════════════════════════════════

def make_ct02_tam_tru():
    doc = new_doc()
    header_vn(doc)
    title(doc, 'TỜ KHAI ĐĂNG KÝ TẠM TRÚ', '(Mẫu CT02 — Thông tư 55/2021/TT-BCA)')
    line(doc, 'Họ và tên:')
    line(doc, 'Ngày, tháng, năm sinh:')
    line(doc, 'Giới tính:', 15); line(doc, 'Dân tộc:', 15); line(doc, 'Quốc tịch:', 15)
    line(doc, 'Số CCCD / CMND:')
    line(doc, 'Số điện thoại liên hệ:')
    doc.add_paragraph()
    p = doc.add_paragraph(); p.add_run('Đề nghị đăng ký tạm trú tại địa chỉ:').bold = True
    line(doc, 'Số nhà, đường/xóm/thôn:')
    line(doc, 'Xã/phường/thị trấn:')
    line(doc, 'Quận/huyện/thị xã:')
    line(doc, 'Tỉnh/thành phố:')
    doc.add_paragraph()
    line(doc, 'Thời hạn tạm trú từ ngày:'); line(doc, 'đến ngày:')
    line(doc, 'Mục đích tạm trú:')
    doc.add_paragraph()
    p = doc.add_paragraph(); p.add_run('Chủ hộ/người cho ở nhờ:').bold = True
    line(doc, 'Họ và tên chủ hộ:')
    line(doc, 'Số CCCD chủ hộ:')
    p2 = doc.add_paragraph()
    p2.add_run('Chữ ký chủ hộ (xác nhận đồng ý):').bold = True
    p2.add_run('  _______________________')
    sig_block(doc, 'NGƯỜI ĐĂNG KÝ TẠM TRÚ')
    save(doc, 'mau-CT02-to-khai-tam-tru.docx')


def make_hoan_cong():
    doc = new_doc()
    header_vn(doc)
    title(doc, 'THÔNG BÁO HOÀN THÀNH XÂY DỰNG CÔNG TRÌNH',
          '(Mẫu số 06 — Thông tư 15/2016/TT-BXD)')
    p = doc.add_paragraph(); p.add_run('Kính gửi:').bold = True
    p.add_run(' Phòng Quản lý đô thị / Phòng Kinh tế - Hạ tầng .................')
    doc.add_paragraph()
    p = doc.add_paragraph(); p.add_run('I. THÔNG TIN CHỦ ĐẦU TƯ').bold = True
    line(doc, 'Họ và tên / Tên tổ chức:')
    line(doc, 'Số CCCD / Mã số doanh nghiệp:')
    line(doc, 'Địa chỉ liên hệ:')
    line(doc, 'Số điện thoại:', 20)
    doc.add_paragraph()
    p = doc.add_paragraph(); p.add_run('II. THÔNG TIN CÔNG TRÌNH').bold = True
    line(doc, 'Tên công trình:')
    line(doc, 'Địa điểm xây dựng:')
    line(doc, 'Số Giấy phép xây dựng:')
    line(doc, 'Ngày cấp:', 20)
    line(doc, 'Loại công trình:')
    line(doc, 'Cấp công trình:', 20)
    line(doc, 'Diện tích xây dựng (m²):', 15)
    line(doc, 'Số tầng:', 10)
    line(doc, 'Ngày khởi công:', 20)
    line(doc, 'Ngày hoàn thành:', 20)
    doc.add_paragraph()
    p = doc.add_paragraph(); p.add_run('III. CAM KẾT').bold = True
    doc.add_paragraph('Tôi/Chúng tôi cam đoan công trình đã được xây dựng đúng thiết kế được duyệt, '
                      'đúng Giấy phép xây dựng được cấp và tuân thủ các quy định của pháp luật về xây dựng.')
    sig_block(doc, 'CHỦ ĐẦU TƯ')
    save(doc, 'mau-thong-bao-hoan-cong-cong-trinh.docx')


def make_chung_thuc():
    doc = new_doc()
    header_vn(doc)
    title(doc, 'YÊU CẦU CHỨNG THỰC BẢN SAO TỪ BẢN CHÍNH',
          '(Theo Nghị định 23/2015/NĐ-CP)')
    p = doc.add_paragraph(); p.add_run('Kính gửi:').bold = True
    p.add_run(' UBND xã/phường/thị trấn ..................................................')
    doc.add_paragraph()
    p = doc.add_paragraph(); p.add_run('I. THÔNG TIN NGƯỜI YÊU CẦU').bold = True
    line(doc, 'Họ và tên:')
    line(doc, 'Ngày sinh:', 20)
    line(doc, 'Số CCCD / CMND:')
    line(doc, 'Địa chỉ thường trú:')
    line(doc, 'Số điện thoại:', 20)
    doc.add_paragraph()
    p = doc.add_paragraph(); p.add_run('II. NỘI DUNG YÊU CẦU CHỨNG THỰC').bold = True
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i, txt in enumerate(['STT', 'Tên giấy tờ cần chứng thực', 'Số trang', 'Số bản sao']):
        hdr[i].text = txt
        hdr[i].paragraphs[0].runs[0].bold = True
    for i in range(5):
        row = tbl.add_row().cells
        row[0].text = str(i + 1)
    doc.add_paragraph()
    line(doc, 'Mục đích sử dụng bản sao:')
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Tôi cam đoan bản chính xuất trình là thật, nội dung hoàn toàn trung thực. '
              'Chịu trách nhiệm pháp lý về lời cam đoan này.')
    sig_block(doc, 'NGƯỜI YÊU CẦU')
    save(doc, 'mau-don-yeu-cau-chung-thuc-ban-sao.docx')


def make_tam_ngung_kd():
    doc = new_doc()
    header_vn(doc)
    title(doc, 'THÔNG BÁO TẠM NGỪNG KINH DOANH / TIẾP TỤC KINH DOANH',
          '(Mẫu II-19 — Nghị định 01/2021/NĐ-CP)')
    p = doc.add_paragraph(); p.add_run('Kính gửi:').bold = True
    p.add_run(' Phòng Đăng ký kinh doanh — Sở Kế hoạch và Đầu tư tỉnh/thành phố .........')
    doc.add_paragraph()
    p = doc.add_paragraph(); p.add_run('I. THÔNG TIN DOANH NGHIỆP / HỘ KINH DOANH').bold = True
    line(doc, 'Tên doanh nghiệp / hộ KD:')
    line(doc, 'Mã số doanh nghiệp / hộ KD:')
    line(doc, 'Địa chỉ trụ sở:')
    line(doc, 'Người đại diện pháp luật:')
    line(doc, 'Chức danh:', 20)
    line(doc, 'Số CCCD:', 20)
    line(doc, 'Số điện thoại:', 20)
    doc.add_paragraph()
    p = doc.add_paragraph(); p.add_run('II. NỘI DUNG THÔNG BÁO').bold = True
    for opt in ['□  Tạm ngừng kinh doanh', '□  Tiếp tục kinh doanh trước thời hạn', '□  Tiếp tục kinh doanh đúng hạn']:
        doc.add_paragraph(opt)
    line(doc, 'Thời gian tạm ngừng từ ngày:'); line(doc, 'đến ngày:')
    line(doc, 'Lý do tạm ngừng:')
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Doanh nghiệp / hộ KD cam kết thực hiện đầy đủ nghĩa vụ tài chính với Nhà nước trong thời gian tạm ngừng kinh doanh.').italic = True
    sig_block(doc, 'ĐẠI DIỆN DOANH NGHIỆP / HỘ KD')
    save(doc, 'mau-thong-bao-tam-ngung-kinh-doanh.docx')


def make_nuoi_con_nuoi():
    doc = new_doc()
    header_vn(doc)
    title(doc, 'TỜ KHAI ĐĂNG KÝ NUÔI CON NUÔI TRONG NƯỚC',
          '(Mẫu TP/CCN-2014-TKNNN.1)')
    p = doc.add_paragraph(); p.add_run('PHẦN I — THÔNG TIN CHA/MẸ NUÔI').bold = True
    line(doc, 'Họ và tên cha nuôi:')
    line(doc, 'Ngày sinh:', 20); line(doc, 'Dân tộc:', 15); line(doc, 'Quốc tịch:', 15)
    line(doc, 'Số CCCD / CMND cha nuôi:')
    line(doc, 'Họ và tên mẹ nuôi:')
    line(doc, 'Ngày sinh mẹ nuôi:', 20); line(doc, 'Số CCCD mẹ nuôi:')
    line(doc, 'Địa chỉ thường trú:')
    line(doc, 'Nghề nghiệp / thu nhập ổn định:', 25)
    doc.add_paragraph()
    p = doc.add_paragraph(); p.add_run('PHẦN II — THÔNG TIN NGƯỜI ĐƯỢC NHẬN LÀM CON NUÔI').bold = True
    line(doc, 'Họ và tên con nuôi:')
    line(doc, 'Ngày, tháng, năm sinh:')
    line(doc, 'Giới tính:', 10); line(doc, 'Dân tộc:', 15)
    line(doc, 'Nơi sinh:')
    line(doc, 'Đang sống cùng (cha/mẹ đẻ / cơ sở nuôi dưỡng):')
    doc.add_paragraph()
    p = doc.add_paragraph(); p.add_run('PHẦN III — CAM KẾT').bold = True
    doc.add_paragraph('Chúng tôi tự nguyện nhận con nuôi và cam kết chăm sóc, nuôi dưỡng, '
                      'giáo dục con nuôi như con đẻ, bảo đảm mọi quyền lợi hợp pháp cho con nuôi.')
    sig_block(doc, 'CHA/MẸ NUÔI')
    save(doc, 'mau-to-khai-dang-ky-nuoi-con-nuoi.docx')


def make_d02_lt_bhxh():
    doc = new_doc()
    header_vn(doc)
    title(doc, 'DANH SÁCH LAO ĐỘNG THAM GIA BHXH, BHYT, BHTN',
          '(Mẫu D02-LT — Ban hành kèm theo QĐ 595/QĐ-BHXH)')
    p = doc.add_paragraph(); p.add_run('Đơn vị:').bold = True; p.add_run(' ' + '_' * 35)
    p = doc.add_paragraph(); p.add_run('Mã đơn vị BHXH:').bold = True; p.add_run(' ' + '_' * 20)
    p = doc.add_paragraph(); p.add_run('Tháng/năm:').bold = True; p.add_run(' ' + '_' * 15)
    doc.add_paragraph()
    tbl = doc.add_table(rows=1, cols=7)
    tbl.style = 'Table Grid'
    hdrs = ['STT', 'Họ và tên', 'Mã số BHXH', 'Ngày sinh', 'Giới tính', 'Lương đóng BH', 'Ghi chú']
    for i, h in enumerate(hdrs):
        tbl.rows[0].cells[i].text = h
        tbl.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        tbl.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(9)
    for i in range(10):
        row = tbl.add_row().cells
        row[0].text = str(i + 1)
        for c in row[1:]:
            if c.paragraphs[0].runs:
                c.paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Tổng số lao động:').bold = True; p.add_run(' ________  người')
    p = doc.add_paragraph()
    p.add_run('Tổng quỹ lương đóng BH:').bold = True; p.add_run(' ________________  đồng')
    sig_block(doc, 'ĐẠI DIỆN ĐƠN VỊ\n(Ký, đóng dấu)')
    save(doc, 'mau-D02-LT-danh-sach-lao-dong-bhxh.docx')


def make_tk1_ts_bhxh():
    doc = new_doc()
    header_vn(doc)
    title(doc, 'TỜ KHAI THAM GIA BHXH, BHYT (Mẫu TK1-TS)',
          '(Dành cho người lao động — QĐ 595/QĐ-BHXH)')
    p = doc.add_paragraph(); p.add_run('PHẦN I — THÔNG TIN CÁ NHÂN').bold = True
    line(doc, 'Họ và tên:')
    line(doc, 'Ngày sinh:'); line(doc, 'Giới tính:', 10); line(doc, 'Dân tộc:', 15)
    line(doc, 'Quốc tịch:', 15)
    line(doc, 'Số CCCD / CMND:')
    line(doc, 'Nơi đăng ký khai sinh:')
    line(doc, 'Địa chỉ thường trú:')
    line(doc, 'Địa chỉ nơi ở hiện tại:')
    line(doc, 'Số điện thoại:', 20); line(doc, 'Email:', 25)
    doc.add_paragraph()
    p = doc.add_paragraph(); p.add_run('PHẦN II — THÔNG TIN BHXH / BHYT').bold = True
    line(doc, 'Mã số BHXH (nếu đã có):')
    line(doc, 'Số sổ BHXH (nếu đã có):')
    line(doc, 'Nơi đăng ký KCB ban đầu:')
    doc.add_paragraph()
    for opt in ['□  Tham gia lần đầu', '□  Điều chỉnh thông tin', '□  Cấp lại sổ / thẻ']:
        doc.add_paragraph(opt)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Tôi cam đoan các thông tin trên là trung thực và chịu trách nhiệm trước pháp luật.').italic = True
    sig_block(doc, 'NGƯỜI THAM GIA')
    save(doc, 'mau-TK1-TS-to-khai-tham-gia-bhxh-bhyt.docx')


def make_to_khai_dang_ky_xe_may():
    doc = new_doc()
    header_vn(doc)
    title(doc, 'GIẤY KHAI ĐĂNG KÝ XE MÔ TÔ / XE MÁY LẦN ĐẦU',
          '(Mẫu 01 — Thông tư 58/2020/TT-BCA)')
    p = doc.add_paragraph(); p.add_run('Kính gửi:').bold = True
    p.add_run(' Phòng Cảnh sát giao thông / Công an .........................................')
    doc.add_paragraph()
    p = doc.add_paragraph(); p.add_run('I. THÔNG TIN CHỦ XE').bold = True
    line(doc, 'Họ và tên:')
    line(doc, 'Ngày sinh:', 20); line(doc, 'Giới tính:', 10)
    line(doc, 'Số CCCD / CMND:')
    line(doc, 'Địa chỉ thường trú:')
    line(doc, 'Số điện thoại:', 20)
    doc.add_paragraph()
    p = doc.add_paragraph(); p.add_run('II. THÔNG TIN XE').bold = True
    line(doc, 'Nhãn hiệu / Model xe:')
    line(doc, 'Màu sơn:', 20)
    line(doc, 'Số khung:'); line(doc, 'Số máy:')
    line(doc, 'Dung tích xi lanh (cc):', 15)
    line(doc, 'Năm sản xuất:', 15)
    line(doc, 'Nước sản xuất:', 15)
    line(doc, 'Nguồn gốc:')
    for opt in ['□  Mua mới từ đại lý', '□  Mua lại từ cá nhân', '□  Nhập khẩu', '□  Được tặng/thừa kế']:
        doc.add_paragraph(opt)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Đề nghị Cơ quan đăng ký xe xem xét, đăng ký và cấp Giấy chứng nhận đăng ký xe, biển số xe.').italic = True
    sig_block(doc, 'CHỦ XE')
    save(doc, 'mau-to-khai-dang-ky-xe-may-lan-dau.docx')


# ════════════════════════════════════════════════════════════════════════════════
# PDF TEMPLATES
# ════════════════════════════════════════════════════════════════════════════════

def make_pdf_bien_ban_ban_giao_nha():
    pdf = VnPDF()
    pdf.add_page()
    pdf.vheader('BIÊN BẢN BÀN GIAO NHÀ Ở / ĐẤT Ở', '(Dùng kèm hợp đồng mua bán / chuyển nhượng)')
    pdf.para('Hôm nay, ngày ...... tháng ...... năm .........., tại địa chỉ: ................................')
    pdf.ln(2)
    pdf.para('BÊN GIAO (Bên A):', bold=True)
    pdf.field('Họ và tên:')
    pdf.field('Số CCCD/CMND:')
    pdf.field('Địa chỉ thường trú:')
    pdf.ln(3)
    pdf.para('BÊN NHẬN (Bên B):', bold=True)
    pdf.field('Họ và tên:')
    pdf.field('Số CCCD/CMND:')
    pdf.field('Địa chỉ thường trú:')
    pdf.ln(3)
    pdf.para('NỘI DUNG BÀN GIAO:', bold=True)
    pdf.field('Địa chỉ tài sản bàn giao:')
    pdf.field('Diện tích (m²):')
    pdf.field('Số tờ bản đồ:'); pdf.field('Số thửa đất:')
    pdf.field('Số GCN QSDĐ (Sổ đỏ/Sổ hồng):')
    pdf.ln(3)
    pdf.para('TÌNH TRẠNG TÀI SẢN KHI BÀN GIAO:', bold=True)
    pdf.multi_cell(0, 6, '_' * 80 + '\n' + '_' * 80 + '\n' + '_' * 80)
    pdf.ln(3)
    pdf.para('Bên A đã giao và Bên B đã nhận đầy đủ theo nội dung trên. Biên bản được lập thành 02 bản, '
             'mỗi bên giữ 01 bản có giá trị pháp lý như nhau.')
    pdf.ln(6)
    pdf.set_font('Arial', 'B', 11)
    pdf.cell(95, 6, 'BÊN GIAO (Bên A)', align='C')
    pdf.cell(95, 6, 'BÊN NHẬN (Bên B)', align='C', new_x='LMARGIN', new_y='NEXT')
    pdf.set_font('Arial', '', 10)
    pdf.cell(95, 6, '(Ký, ghi rõ họ tên)', align='C')
    pdf.cell(95, 6, '(Ký, ghi rõ họ tên)', align='C', new_x='LMARGIN', new_y='NEXT')
    save_pdf(pdf, 'mau-bien-ban-ban-giao-nha-dat.pdf')


def make_pdf_hop_dong_thue_nha():
    pdf = VnPDF()
    pdf.add_page()
    pdf.vheader('HỢP ĐỒNG THUÊ NHÀ Ở', '(Theo Luật Nhà ở 2023)')
    pdf.para('Hôm nay, ngày ...... tháng ...... năm ......, tại ............................................')
    pdf.para('Hai bên gồm:', bold=True)
    pdf.ln(2)
    pdf.para('BÊN CHO THUÊ (Bên A):', bold=True)
    pdf.field('Họ và tên:'); pdf.field('Số CCCD:')
    pdf.field('Địa chỉ liên hệ:'); pdf.field('Số điện thoại:')
    pdf.ln(3)
    pdf.para('BÊN THUÊ (Bên B):', bold=True)
    pdf.field('Họ và tên:'); pdf.field('Số CCCD:')
    pdf.field('Địa chỉ liên hệ:'); pdf.field('Số điện thoại:')
    pdf.ln(3)
    pdf.para('ĐIỀU 1 — ĐỐI TƯỢNG HỢP ĐỒNG', bold=True)
    pdf.field('Địa chỉ nhà thuê:'); pdf.field('Diện tích sử dụng (m²):')
    pdf.field('Mục đích sử dụng:')
    pdf.ln(3)
    pdf.para('ĐIỀU 2 — THỜI HẠN THUÊ', bold=True)
    pdf.field('Từ ngày:'); pdf.field('Đến ngày:')
    pdf.ln(3)
    pdf.para('ĐIỀU 3 — GIÁ THUÊ VÀ PHƯƠNG THỨC THANH TOÁN', bold=True)
    pdf.field('Giá thuê hàng tháng:')
    pdf.field('Đặt cọc:')
    pdf.field('Thanh toán vào ngày:'); pdf.field('hàng tháng')
    pdf.ln(3)
    pdf.para('ĐIỀU 4 — QUYỀN VÀ NGHĨA VỤ CÁC BÊN', bold=True)
    pdf.multi_cell(0, 6, 'Theo quy định của Luật Nhà ở và các thỏa thuận cụ thể đính kèm.')
    pdf.ln(6)
    pdf.set_font('Arial', 'B', 11)
    pdf.cell(95, 6, 'BÊN CHO THUÊ (Bên A)', align='C')
    pdf.cell(95, 6, 'BÊN THUÊ (Bên B)', align='C', new_x='LMARGIN', new_y='NEXT')
    pdf.set_font('Arial', '', 10)
    pdf.cell(95, 6, '(Ký, ghi rõ họ tên)', align='C')
    pdf.cell(95, 6, '(Ký, ghi rõ họ tên)', align='C', new_x='LMARGIN', new_y='NEXT')
    save_pdf(pdf, 'mau-hop-dong-thue-nha.pdf')


def make_pdf_cam_ket_khong_tranh_chap():
    pdf = VnPDF()
    pdf.add_page()
    pdf.vheader('GIẤY CAM KẾT KHÔNG TRANH CHẤP ĐẤT ĐAI', '(Dùng khi làm thủ tục đăng ký đất đai)')
    pdf.para('Kính gửi: UBND xã/phường/thị trấn .........................................................')
    pdf.ln(4)
    pdf.para('Tôi (chúng tôi) là:', bold=True)
    pdf.field('Họ và tên:'); pdf.field('Ngày sinh:')
    pdf.field('Số CCCD / CMND:'); pdf.field('Địa chỉ thường trú:')
    pdf.ln(3)
    pdf.para('Là chủ sử dụng / đề nghị đăng ký quyền sử dụng đối với thửa đất:', bold=True)
    pdf.field('Địa chỉ thửa đất:'); pdf.field('Số thửa:'); pdf.field('Số tờ bản đồ:')
    pdf.field('Diện tích (m²):'); pdf.field('Mục đích sử dụng:')
    pdf.ln(3)
    pdf.para('NỘI DUNG CAM KẾT:', bold=True)
    for i, item in enumerate([
        'Thửa đất nêu trên thuộc quyền sử dụng hợp pháp của tôi/chúng tôi.',
        'Không có tranh chấp, khiếu nại, khởi kiện về quyền sử dụng thửa đất.',
        'Nguồn gốc đất rõ ràng, không có vi phạm pháp luật đất đai.',
        'Nếu có tranh chấp phát sinh, tôi/chúng tôi hoàn toàn chịu trách nhiệm trước pháp luật.',
    ], 1):
        pdf.para(f'{i}. {item}')
    pdf.sig('NGƯỜI CAM KẾT')
    pdf.ln(6)
    pdf.para('XÁC NHẬN CỦA UBND XÃ/PHƯỜNG/THỊ TRẤN:', bold=True)
    pdf.multi_cell(0, 6, '_' * 80 + '\n' + '_' * 80)
    save_pdf(pdf, 'mau-giay-cam-ket-khong-tranh-chap-dat.pdf')


def make_pdf_bang_ke_ho_so():
    pdf = VnPDF()
    pdf.add_page()
    pdf.vheader('BẢNG KÊ HỒ SƠ NỘP TẠI BỘ PHẬN MỘT CỬA', '(Phiếu tiếp nhận hồ sơ — lưu tại cơ quan)')
    pdf.field('Số tiếp nhận:'); pdf.field('Ngày tiếp nhận:')
    pdf.field('Thủ tục hành chính:')
    pdf.field('Cơ quan giải quyết:')
    pdf.ln(3)
    pdf.para('THÔNG TIN NGƯỜI NỘP HỒ SƠ:', bold=True)
    pdf.field('Họ và tên:'); pdf.field('Số CCCD/CMND:')
    pdf.field('Địa chỉ:'); pdf.field('Số điện thoại:')
    pdf.ln(3)
    pdf.para('DANH MỤC GIẤY TỜ NỘP:', bold=True)
    # Table header
    pdf.set_font('Arial', 'B', 10)
    pdf.cell(10, 7, 'STT', border=1, align='C')
    pdf.cell(100, 7, 'Tên giấy tờ', border=1, align='C')
    pdf.cell(25, 7, 'Số lượng', border=1, align='C')
    pdf.cell(25, 7, 'Bản chính', border=1, align='C')
    pdf.cell(25, 7, 'Bản sao', border=1, align='C', new_x='LMARGIN', new_y='NEXT')
    pdf.set_font('Arial', '', 10)
    for i in range(8):
        pdf.cell(10, 7, str(i+1), border=1, align='C')
        pdf.cell(100, 7, '', border=1)
        pdf.cell(25, 7, '', border=1, align='C')
        pdf.cell(25, 7, '', border=1, align='C')
        pdf.cell(25, 7, '', border=1, align='C', new_x='LMARGIN', new_y='NEXT')
    pdf.ln(3)
    pdf.field('Hẹn trả kết quả ngày:'); pdf.field('Lệ phí:')
    pdf.ln(6)
    pdf.set_font('Arial', 'B', 11)
    pdf.cell(95, 6, 'NGƯỜI NỘP HỒ SƠ', align='C')
    pdf.cell(95, 6, 'CÔNG CHỨC TIẾP NHẬN', align='C', new_x='LMARGIN', new_y='NEXT')
    pdf.set_font('Arial', '', 10)
    pdf.cell(95, 6, '(Ký, ghi rõ họ tên)', align='C')
    pdf.cell(95, 6, '(Ký, đóng dấu)', align='C', new_x='LMARGIN', new_y='NEXT')
    save_pdf(pdf, 'mau-bang-ke-ho-so-nop-bo-phan-mot-cua.pdf')


def make_pdf_giay_xac_nhan_cu_tru():
    pdf = VnPDF()
    pdf.add_page()
    pdf.vheader('GIẤY XÁC NHẬN THÔNG TIN CƯ TRÚ', '(Theo Luật Cư trú 2020 — Thông tư 55/2021/TT-BCA)')
    pdf.para('Kính gửi: Cơ quan / Tổ chức có yêu cầu xác nhận:')
    pdf.multi_cell(0, 6, '...............................................................................')
    pdf.ln(3)
    pdf.para('UBND xã/phường/thị trấn .................... xác nhận:', bold=True)
    pdf.ln(3)
    pdf.field('Họ và tên:'); pdf.field('Ngày sinh:')
    pdf.field('Giới tính:'); pdf.field('Dân tộc:')
    pdf.field('Số CCCD / CMND:'); pdf.field('Quốc tịch:')
    pdf.ln(2)
    pdf.para('THÔNG TIN CƯ TRÚ:', bold=True)
    pdf.field('Loại cư trú:')
    for opt in ['□  Thường trú', '□  Tạm trú', '□  Tạm vắng']:
        pdf.set_font('Arial', '', 11)
        pdf.cell(0, 6, opt, new_x='LMARGIN', new_y='NEXT')
    pdf.field('Địa chỉ đăng ký:'); pdf.field('Từ ngày:')
    pdf.field('Chủ hộ:'); pdf.field('Quan hệ với chủ hộ:')
    pdf.ln(3)
    pdf.para('Xác nhận để sử dụng vào mục đích:')
    pdf.multi_cell(0, 6, '...............................................................................')
    pdf.ln(6)
    pdf.set_font('Arial', 'B', 11)
    pdf.cell(0, 6, 'TM. UBND XÃ/PHƯỜNG/THỊ TRẤN', align='R', new_x='LMARGIN', new_y='NEXT')
    pdf.set_font('Arial', '', 10)
    pdf.cell(0, 6, 'CHỦ TỊCH / PHÓ CHỦ TỊCH', align='R', new_x='LMARGIN', new_y='NEXT')
    pdf.cell(0, 6, '(Ký tên, đóng dấu)', align='R', new_x='LMARGIN', new_y='NEXT')
    save_pdf(pdf, 'mau-giay-xac-nhan-cu-tru.pdf')


# ════════════════════════════════════════════════════════════════════════════════
if __name__ == '__main__':
    print('=== Tao DOCX templates ===')
    make_ct02_tam_tru()
    make_hoan_cong()
    make_chung_thuc()
    make_tam_ngung_kd()
    make_nuoi_con_nuoi()
    make_d02_lt_bhxh()
    make_tk1_ts_bhxh()
    make_to_khai_dang_ky_xe_may()

    print('\n=== Tao PDF templates ===')
    make_pdf_bien_ban_ban_giao_nha()
    make_pdf_hop_dong_thue_nha()
    make_pdf_cam_ket_khong_tranh_chap()
    make_pdf_bang_ke_ho_so()
    make_pdf_giay_xac_nhan_cu_tru()

    # Tong ket
    import os
    tdir = os.path.normpath(os.path.join(os.path.dirname(__file__), '..', 'data', 'templates'))
    files = [f for f in os.listdir(tdir) if f.endswith(('.doc', '.docx', '.pdf'))]
    print(f'\nTong files: {len(files)} (docx: {sum(1 for f in files if f.endswith(".docx"))}, '
          f'doc: {sum(1 for f in files if f.endswith(".doc"))}, '
          f'pdf: {sum(1 for f in files if f.endswith(".pdf"))})')
