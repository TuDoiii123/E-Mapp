"""
Tao lai cac file .docx bi luu sai dinh dang (HTML content / OLE2 sai extension).
"""
import sys, os, shutil
sys.path.insert(0, '.')

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

TDIR = os.path.normpath(os.path.join(os.path.dirname(__file__), '..', 'data', 'templates'))


def header(doc, line1, line2=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line1); r.bold = True; r.font.size = Pt(12)
    if line2:
        doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(line2); r2.bold = True; r2.font.size = Pt(12)


def title(doc, text, sub=None):
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(13)
    if sub:
        p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.add_run(sub).font.size = Pt(10)
    doc.add_paragraph()


def field(doc, label, width=40):
    p = doc.add_paragraph()
    r = p.add_run(label); r.bold = True; r.font.size = Pt(11)
    p.add_run(' ' + '_' * width).font.size = Pt(11)


def sig(doc, role='NGƯỜI YÊU CẦU'):
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('........, ngày ...... tháng ...... năm ......').italic = True
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(role).bold = True
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('(Ký, ghi rõ họ tên)').italic = True


def margins(doc):
    s = doc.sections[0]
    s.left_margin = Cm(2.5); s.right_margin = Cm(2)
    s.top_margin = Cm(2); s.bottom_margin = Cm(1.5)


# ── 1. mau-to-khai-dang-ky-lai-khai-sinh.docx ────────────────────────────────
def make_lai_khai_sinh():
    doc = Document(); margins(doc)
    header(doc, 'CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM',
                'Độc lập - Tự do - Hạnh phúc')
    title(doc, 'TỜ KHAI ĐĂNG KÝ LẠI KHAI SINH', '(Mẫu HTT-2014-01.3)')

    field(doc, 'Họ và tên người yêu cầu:')
    field(doc, 'Ngày, tháng, năm sinh:')
    field(doc, 'Giới tính:', 15)
    field(doc, 'Dân tộc:', 15)
    field(doc, 'Quốc tịch:', 15)
    field(doc, 'Số CCCD / CMND:')
    field(doc, 'Địa chỉ thường trú:')
    field(doc, 'Số điện thoại:', 20)
    doc.add_paragraph()

    p = doc.add_paragraph(); p.add_run('Đề nghị đăng ký lại khai sinh cho:').bold = True
    field(doc, 'Họ và tên người được đăng ký lại:')
    field(doc, 'Ngày, tháng, năm sinh:')
    field(doc, 'Giới tính:', 15)
    field(doc, 'Dân tộc:', 15)
    field(doc, 'Quốc tịch:', 15)
    field(doc, 'Nơi sinh:')
    field(doc, 'Họ và tên cha:')
    field(doc, 'Ngày sinh cha:', 20)
    field(doc, 'Họ và tên mẹ:')
    field(doc, 'Ngày sinh mẹ:', 20)
    doc.add_paragraph()

    p = doc.add_paragraph(); p.add_run('Lý do đăng ký lại:').bold = True
    p.add_run(' ' + '_' * 45)
    field(doc, 'Tên cơ quan đã đăng ký khai sinh trước đây (nếu có):')

    sig(doc)
    path = os.path.join(TDIR, 'mau-to-khai-dang-ky-lai-khai-sinh.docx')
    doc.save(path)
    print(f'  OK  mau-to-khai-dang-ky-lai-khai-sinh.docx  ({os.path.getsize(path)//1024}KB)')


# ── 2. mau-don-tach-thua-hop-thua-dat.docx ───────────────────────────────────
def make_tach_thua_dat():
    doc = Document(); margins(doc)
    header(doc, 'CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM',
                'Độc lập - Tự do - Hạnh phúc')
    title(doc, 'ĐƠN ĐỀ NGHỊ TÁCH THỬA / HỢP THỬA ĐẤT', '(Mẫu 11/ĐĐ — Thông tư 24/2014/TT-BTNMT)')

    p = doc.add_paragraph(); p.add_run('Kính gửi:').bold = True
    p.add_run(' Văn phòng Đăng ký đất đai / Chi nhánh Văn phòng Đăng ký đất đai')

    doc.add_paragraph()
    p = doc.add_paragraph(); p.add_run('I. THÔNG TIN NGƯỜI SỬ DỤNG ĐẤT').bold = True
    field(doc, 'Họ và tên (hộ gia đình / cá nhân / tổ chức):')
    field(doc, 'Năm sinh:', 20)
    field(doc, 'Số CCCD / CMND / Hộ chiếu:')
    field(doc, 'Địa chỉ thường trú / trụ sở:')
    field(doc, 'Số điện thoại:', 20)

    doc.add_paragraph()
    p = doc.add_paragraph(); p.add_run('II. THÔNG TIN THỬA ĐẤT').bold = True
    field(doc, 'Thửa đất số:')
    field(doc, 'Tờ bản đồ số:')
    field(doc, 'Địa chỉ thửa đất:')
    field(doc, 'Diện tích hiện tại (m²):', 20)
    field(doc, 'Mục đích sử dụng:')
    field(doc, 'Số GCN QSDĐ (Sổ đỏ/Sổ hồng):')

    doc.add_paragraph()
    p = doc.add_paragraph(); p.add_run('III. NỘI DUNG ĐỀ NGHỊ').bold = True
    for opt in ['□  Tách thửa đất', '□  Hợp thửa đất']:
        doc.add_paragraph(opt)

    field(doc, 'Diện tích thửa sau tách/hợp (dự kiến):')
    p = doc.add_paragraph(); p.add_run('Lý do:').bold = True
    p.add_run(' ' + '_' * 50)

    sig(doc, 'NGƯỜI LÀM ĐƠN')
    path = os.path.join(TDIR, 'mau-don-tach-thua-hop-thua-dat.docx')
    doc.save(path)
    print(f'  OK  mau-don-tach-thua-hop-thua-dat.docx  ({os.path.getsize(path)//1024}KB)')


# ── 3. mau-to-khai-cap-the-bao-hiem-y-te.docx ────────────────────────────────
def make_bhyt():
    doc = Document(); margins(doc)
    header(doc, 'CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM',
                'Độc lập - Tự do - Hạnh phúc')
    title(doc, 'TỜ KHAI CẤP / ĐỔI THẺ BẢO HIỂM Y TẾ',
          '(Mẫu TK3-TS — Ban hành kèm theo QĐ 595/QĐ-BHXH)')

    p = doc.add_paragraph(); p.add_run('PHẦN I — THÔNG TIN NGƯỜI THAM GIA').bold = True
    field(doc, 'Họ và tên:')
    field(doc, 'Ngày sinh:')
    field(doc, 'Giới tính:', 10)
    field(doc, 'Mã số BHXH / Số sổ BHXH:')
    field(doc, 'Số CCCD / CMND:')
    field(doc, 'Địa chỉ thường trú:')
    field(doc, 'Số điện thoại:', 20)
    doc.add_paragraph()

    p = doc.add_paragraph(); p.add_run('PHẦN II — NỘI DUNG ĐỀ NGHỊ').bold = True
    for opt in [
        '□  Cấp thẻ BHYT lần đầu',
        '□  Đổi thẻ BHYT (hỏng, mất, thay đổi thông tin)',
        '□  Cấp lại thẻ BHYT',
    ]:
        doc.add_paragraph(opt)

    doc.add_paragraph()
    p = doc.add_paragraph(); p.add_run('Nơi đăng ký KCB ban đầu:').bold = True
    p.add_run(' ' + '_' * 35)
    field(doc, 'Từ ngày:', 20)
    doc.add_paragraph()

    p = doc.add_paragraph(); p.add_run('PHẦN III — CAM KẾT').bold = True
    p2 = doc.add_paragraph()
    p2.add_run('Tôi cam đoan các thông tin khai trên là đúng sự thật và chịu trách nhiệm '
               'trước pháp luật về những thông tin đã khai.')

    sig(doc, 'NGƯỜI THAM GIA BHYT')
    path = os.path.join(TDIR, 'mau-to-khai-cap-the-bao-hiem-y-te.docx')
    doc.save(path)
    print(f'  OK  mau-to-khai-cap-the-bao-hiem-y-te.docx  ({os.path.getsize(path)//1024}KB)')


# ── 4. mau-to-khai-cap-CMND.docx (khong assign nhung fix luon) ───────────────
def make_cmnd():
    doc = Document(); margins(doc)
    header(doc, 'CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM',
                'Độc lập - Tự do - Hạnh phúc')
    title(doc, 'TỜ KHAI CẤP / ĐỔI / CẤP LẠI CHỨNG MINH NHÂN DÂN',
          '(Mẫu CM01 — Thông tư 06/2013/TT-BCA)')

    field(doc, 'Họ và tên:')
    field(doc, 'Ngày, tháng, năm sinh:')
    field(doc, 'Giới tính:', 10)
    field(doc, 'Dân tộc:', 15)
    field(doc, 'Tôn giáo:', 15)
    field(doc, 'Quê quán:')
    field(doc, 'Địa chỉ thường trú:')
    field(doc, 'Số CMND cũ (nếu có):')
    doc.add_paragraph()

    p = doc.add_paragraph(); p.add_run('Lý do đề nghị:').bold = True
    for opt in ['□  Cấp mới', '□  Đổi (hỏng/mất/hết hạn)', '□  Cấp lại']:
        doc.add_paragraph(opt)

    sig(doc)
    path = os.path.join(TDIR, 'mau-to-khai-cap-CMND.docx')
    doc.save(path)
    print(f'  OK  mau-to-khai-cap-CMND.docx  ({os.path.getsize(path)//1024}KB)')


# ── 5. OLE2 files: doi ten sang .doc ─────────────────────────────────────────
OLE2_RENAMES = [
    ('mau-bien-ban-vi-pham-hanh-chinh.docx', 'mau-bien-ban-vi-pham-hanh-chinh.doc'),
    ('mau-to-khai-cap-ho-chieu.docx',        'mau-to-khai-cap-ho-chieu.doc'),
    ('mau-to-khai-le-phi-truoc-ba.docx',     'mau-to-khai-le-phi-truoc-ba.doc'),
    ('mau-to-khai-thue-thu-nhap-ca-nhan.docx', 'mau-to-khai-thue-thu-nhap-ca-nhan.doc'),
]

def rename_ole2():
    from sqlalchemy import text
    from server import app
    from models.db import db
    renames = []
    with app.app_context():
        for old_name, new_name in OLE2_RENAMES:
            src = os.path.join(TDIR, old_name)
            dst = os.path.join(TDIR, new_name)
            if not os.path.exists(src):
                print(f'  SKIP (not found) {old_name}')
                continue
            if os.path.exists(dst):
                os.remove(src)   # xoa .docx duplicate
                print(f'  SKIP (dst exists, removed src) {old_name}')
                continue
            shutil.move(src, dst)
            renames.append((old_name, new_name))
            print(f'  RENAME  {old_name} -> {new_name}')

        # Cap nhat DB neu co reference
        for old_name, new_name in renames:
            result = db.session.execute(text(
                "UPDATE service_requirements SET template_file = :new WHERE template_file = :old"
            ), {'new': new_name, 'old': old_name})
            if result.rowcount:
                print(f'    DB updated: {result.rowcount} rows: {old_name} -> {new_name}')
        db.session.commit()


if __name__ == '__main__':
    print("=== Tao lai HTML-disguised .docx files ===")
    make_lai_khai_sinh()
    make_tach_thua_dat()
    make_bhyt()
    make_cmnd()

    print("\n=== Doi ten OLE2 .docx -> .doc ===")
    rename_ole2()

    print("\n=== Kiem tra lai ===")
    files = [f for f in os.listdir(TDIR) if f.endswith(('.doc','.docx'))]
    bad = 0
    for fname in sorted(files):
        fpath = os.path.join(TDIR, fname)
        b = open(fpath,'rb').read(4)
        ext = fname.rsplit('.',1)[-1].lower()
        ok = (ext=='docx' and b[:2]==b'\x50\x4b') or (ext=='doc' and b[:2]==b'\xd0\xcf')
        if not ok:
            bad += 1
            print(f'  BAD [{b.hex()}] {fname}')
    print(f"  Files ok: {len(files)-bad}/{len(files)}, bad: {bad}")
