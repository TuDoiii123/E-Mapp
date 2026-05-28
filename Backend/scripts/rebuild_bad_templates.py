#!/usr/bin/env python3
"""
Tạo lại 9 templates đang dùng format bảng 2 cột sai
→ Đúng chuẩn form Việt Nam với dòng kẻ điền thông tin.
"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = os.path.join(os.path.dirname(__file__), '..', 'data', 'templates')


# ── Helpers ───────────────────────────────────────────────────────────────────

def _set_font(run, bold=False, size=12, italic=False, underline=False):
    run.font.name  = 'Times New Roman'
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    run.font.underline = underline
    rPr = run._r.get_or_add_rPr()
    rf  = OxmlElement('w:rFonts')
    rf.set(qn('w:ascii'), 'Times New Roman')
    rf.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.insert(0, rf)


def _add_para(doc, text='', bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
              size=12, sb=0, sa=6, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if text:
        r = p.add_run(text)
        _set_font(r, bold=bold, size=size, italic=italic)
    return p


def _add_field(doc, label, ph, indent=0, sa=4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after  = Pt(sa)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    r1 = p.add_run(label + ': ')
    _set_font(r1, size=12)
    r2 = p.add_run(ph)
    _set_font(r2, size=12, underline=True)
    return p


def _add_two(doc, lbl1, ph1, lbl2, ph2, sa=4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(sa)
    for lbl, ph, sep in [(lbl1, ph1, '        '), (lbl2, ph2, '')]:
        r1 = p.add_run(lbl + ': ')
        _set_font(r1, size=12)
        r2 = p.add_run(ph + sep)
        _set_font(r2, size=12, underline=True)
    return p


def _add_sign_row(doc, left_text='', right_lines=None):
    """Hàng chữ ký 2 cột (có thể bỏ cột trái)."""
    if right_lines is None:
        right_lines = []
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    # Bỏ border
    for row in tbl.rows:
        for cell in row.cells:
            tc  = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBdr = OxmlElement('w:tcBorders')
            for side in ['top', 'bottom', 'left', 'right']:
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'), 'none')
                tcBdr.append(b)
            tcPr.append(tcBdr)
    if left_text:
        p = tbl.cell(0, 0).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(left_text)
        _set_font(r, size=12, italic=True)
    cr = tbl.cell(0, 1)
    for i, line in enumerate(right_lines):
        p = cr.paragraphs[0] if i == 0 else cr.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bold  = line.get('bold', False)
        italic = line.get('italic', False)
        r = p.add_run(line['text'])
        _set_font(r, size=12, bold=bold, italic=italic)


def _header(doc, title, sub=''):
    _add_para(doc, 'CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM',
              bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13, sa=0)
    _add_para(doc, 'Độc lập - Tự do - Hạnh phúc',
              align=WD_ALIGN_PARAGRAPH.CENTER, size=12, sa=2)
    p = doc.add_paragraph('─────────────────────')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(10)
    _add_para(doc, title, bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER, size=14, sa=2 if sub else 14)
    if sub:
        _add_para(doc, sub, align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=11, italic=True, sa=14)


def _new_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(3)
    sec.right_margin  = Cm(2)
    return doc


def _save(doc, filename):
    path = os.path.join(OUT_DIR, filename)
    doc.save(path)
    print(f'✓ {filename}')


# ── 1. Tờ khai đăng ký kết hôn ───────────────────────────────────────────────

def build_ket_hon():
    doc = _new_doc()
    _header(doc, 'TỜ KHAI ĐĂNG KÝ KẾT HÔN',
            '(Ban hành kèm theo Thông tư số 04/2020/TT-BTP ngày 28/5/2020 của Bộ Tư pháp)')
    _add_para(doc, 'Kính gửi: UBND ......................................................................', sa=10)

    _add_para(doc, 'BÊN NAM', bold=True, sa=4)
    _add_field(doc, 'Họ và tên (chữ in hoa)', '{{ho_ten_nguoi_thu_nhat}}')
    _add_two(doc,  'Ngày, tháng, năm sinh', '{{ngay_sinh_nguoi_1}}',
                   'Dân tộc', '{{dan_toc_nguoi_1}}')
    _add_two(doc,  'Quốc tịch', '{{quoc_tich_nguoi_1}}',
                   'Số CCCD/CMND', '{{cccd_nguoi_1}}')
    _add_field(doc, 'Nơi thường trú/tạm trú', '{{dia_chi_nguoi_1}}')

    _add_para(doc, '', sa=4)
    _add_para(doc, 'BÊN NỮ', bold=True, sa=4)
    _add_field(doc, 'Họ và tên (chữ in hoa)', '{{ho_ten_nguoi_thu_hai}}')
    _add_two(doc,  'Ngày, tháng, năm sinh', '{{ngay_sinh_nguoi_2}}',
                   'Dân tộc', '{{dan_toc_nguoi_2}}')
    _add_two(doc,  'Quốc tịch', '{{quoc_tich_nguoi_2}}',
                   'Số CCCD/CMND', '{{cccd_nguoi_2}}')
    _add_field(doc, 'Nơi thường trú/tạm trú', '{{dia_chi_nguoi_2}}')

    _add_para(doc, '', sa=6)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(
        'Hai bên chúng tôi tự nguyện đến đăng ký kết hôn và cam đoan những nội dung '
        'khai trên đây là đúng sự thật, chịu trách nhiệm trước pháp luật về tính '
        'chính xác của các thông tin đã khai.'
    )
    _set_font(r, size=12)

    _add_para(doc, '', sa=6)
    _add_sign_row(doc,
        left_text='Người thứ nhất\n(Ký, ghi rõ họ tên)\n\n\n\n{{ho_ten_nguoi_thu_nhat}}',
        right_lines=[
            {'text': '{{dia_diem}}, ngày {{ngay}} tháng {{thang}} năm {{nam}}', 'italic': True},
            {'text': 'Người thứ hai', 'bold': True},
            {'text': '(Ký, ghi rõ họ tên)', 'italic': True},
            {'text': ''},{'text': ''},{'text': ''},
            {'text': '{{ho_ten_nguoi_thu_hai}}', 'bold': True},
        ]
    )
    _save(doc, 'mau-to-khai-dang-ky-ket-hon.docx')


# ── 2. Tờ khai đăng ký khai tử ───────────────────────────────────────────────

def build_khai_tu():
    doc = _new_doc()
    _header(doc, 'TỜ KHAI ĐĂNG KÝ KHAI TỬ',
            '(Ban hành kèm theo Thông tư số 04/2020/TT-BTP ngày 28/5/2020 của Bộ Tư pháp)')
    _add_para(doc, 'Kính gửi: UBND ......................................................................', sa=10)

    _add_para(doc, 'I. THÔNG TIN NGƯỜI KHAI', bold=True, sa=4)
    _add_field(doc, 'Họ và tên người khai', '{{ho_ten_nguoi_khai}}')
    _add_two(doc, 'Số CCCD/CMND', '{{cccd_nguoi_khai}}',
                  'Quan hệ với người chết', '{{quan_he}}')

    _add_para(doc, '', sa=4)
    _add_para(doc, 'II. THÔNG TIN NGƯỜI ĐÃ CHẾT', bold=True, sa=4)
    _add_field(doc, 'Họ và tên (chữ in hoa)', '{{ho_ten_nguoi_chet}}')
    _add_two(doc, 'Ngày, tháng, năm sinh', '{{ngay_sinh_nguoi_chet}}',
                  'Giới tính', '{{gioi_tinh}}')
    _add_two(doc, 'Dân tộc', '{{dan_toc}}', 'Quốc tịch', '{{quoc_tich}}')
    _add_field(doc, 'Nơi thường trú/tạm trú', '{{dia_chi_thuong_tru}}')
    _add_field(doc, 'Ngày chết (theo giấy chứng tử)', '{{ngay_chet}}')
    _add_field(doc, 'Nơi chết', '{{noi_chet}}')
    _add_field(doc, 'Nguyên nhân chết', '{{nguyen_nhan_chet}}')

    _add_para(doc, '', sa=6)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(
        'Tôi cam đoan những nội dung khai trên đây là đúng sự thật và chịu '
        'trách nhiệm trước pháp luật về tính chính xác của các thông tin đã khai.'
    )
    _set_font(r, size=12)

    _add_para(doc, '', sa=6)
    _add_sign_row(doc, right_lines=[
        {'text': '{{dia_diem}}, ngày {{ngay}} tháng {{thang}} năm {{nam}}', 'italic': True},
        {'text': 'Người khai', 'bold': True},
        {'text': '(Ký, ghi rõ họ tên)', 'italic': True},
        {'text': ''},{'text': ''},{'text': ''},
        {'text': '{{ho_ten_nguoi_khai}}', 'bold': True},
    ])
    _save(doc, 'mau-to-khai-dang-ky-khai-tu.docx')


# ── 3. Tờ khai đăng ký thường trú ────────────────────────────────────────────

def build_thuong_tru():
    doc = _new_doc()
    _header(doc, 'TỜ KHAI ĐĂNG KÝ THƯỜNG TRÚ',
            '(Mẫu HK01 - Ban hành kèm theo Thông tư số 56/2021/TT-BCA ngày 15/5/2021 của Bộ Công an)')
    _add_para(doc, 'Kính gửi: Cơ quan Công an ...................................................', sa=10)

    _add_para(doc, 'I. THÔNG TIN NGƯỜI ĐỀ NGHỊ', bold=True, sa=4)
    _add_field(doc, 'Họ và tên (chữ in hoa)', '{{ho_ten}}')
    _add_two(doc, 'Ngày, tháng, năm sinh', '{{ngay_sinh}}', 'Giới tính', '{{gioi_tinh}}')
    _add_two(doc, 'Dân tộc', '{{dan_toc}}', 'Số CCCD/CMND', '{{cccd}}')
    _add_field(doc, 'Địa chỉ nơi ở hiện tại', '{{dia_chi_hien_tai}}')

    _add_para(doc, '', sa=4)
    _add_para(doc, 'II. THÔNG TIN ĐĂNG KÝ', bold=True, sa=4)
    _add_field(doc, 'Địa chỉ đăng ký thường trú mới', '{{dia_chi_thuong_tru_moi}}')
    _add_field(doc, 'Chủ hộ và quan hệ với chủ hộ', '{{chu_ho_va_quan_he}}')
    _add_field(doc, 'Lý do thay đổi nơi thường trú', '{{ly_do_thay_doi}}')

    _add_para(doc, '', sa=6)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(
        'Tôi cam đoan những nội dung khai trên đây là đúng sự thật và chịu '
        'trách nhiệm trước pháp luật về tính chính xác của các thông tin đã khai.'
    )
    _set_font(r, size=12)

    _add_para(doc, '', sa=6)
    _add_sign_row(doc, right_lines=[
        {'text': '{{dia_diem}}, ngày {{ngay}} tháng {{thang}} năm {{nam}}', 'italic': True},
        {'text': 'Người khai', 'bold': True},
        {'text': '(Ký, ghi rõ họ tên)', 'italic': True},
        {'text': ''},{'text': ''},{'text': ''},
        {'text': '{{ho_ten}}', 'bold': True},
    ])
    _save(doc, 'mau-to-khai-dang-ky-thuong-tru.docx')


# ── 4. Tờ khai đăng ký thuế ──────────────────────────────────────────────────

def build_dang_ky_thue():
    doc = _new_doc()
    _header(doc, 'TỜ KHAI ĐĂNG KÝ THUẾ',
            '(Mẫu 05-ĐK-TCT kèm theo Thông tư số 105/2020/TT-BTC ngày 3/12/2020 của Bộ Tài chính)')
    _add_para(doc, 'Kính gửi: Cơ quan Thuế .....................................................', sa=10)

    _add_para(doc, 'I. THÔNG TIN CÁ NHÂN', bold=True, sa=4)
    _add_field(doc, 'Họ và tên (chữ in hoa)', '{{ho_ten}}')
    _add_two(doc, 'Ngày, tháng, năm sinh', '{{ngay_sinh}}', 'Giới tính', '{{gioi_tinh}}')
    _add_two(doc, 'Quốc tịch', '{{quoc_tich}}', 'Số CCCD/CMND/Hộ chiếu', '{{cccd}}')
    _add_field(doc, 'Địa chỉ thường trú', '{{dia_chi_thuong_tru}}')
    _add_field(doc, 'Địa chỉ hiện tại', '{{dia_chi_hien_tai}}')
    _add_two(doc, 'Số điện thoại', '{{so_dien_thoai}}', 'Email', '{{email}}')

    _add_para(doc, '', sa=4)
    _add_para(doc, 'II. THÔNG TIN THU NHẬP', bold=True, sa=4)
    _add_field(doc, 'Loại thu nhập đăng ký', '{{loai_thu_nhap}}')
    _add_field(doc, 'Nguồn thu nhập chính', '{{nguon_thu_nhap}}')
    _add_field(doc, 'Mã số thuế cũ (nếu có)', '{{ma_so_thue_cu}}')

    _add_para(doc, '', sa=6)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(
        'Tôi cam đoan những nội dung khai trên đây là đúng sự thật và chịu '
        'trách nhiệm trước pháp luật về tính chính xác của các thông tin đã khai.'
    )
    _set_font(r, size=12)

    _add_para(doc, '', sa=6)
    _add_sign_row(doc, right_lines=[
        {'text': '{{dia_diem}}, ngày {{ngay}} tháng {{thang}} năm {{nam}}', 'italic': True},
        {'text': 'Người khai', 'bold': True},
        {'text': '(Ký, ghi rõ họ tên)', 'italic': True},
        {'text': ''},{'text': ''},{'text': ''},
        {'text': '{{ho_ten}}', 'bold': True},
    ])
    _save(doc, 'mau-to-khai-dang-ky-thue.docx')


# ── 5. Đơn đề nghị cấp/đổi GPLX ─────────────────────────────────────────────

def build_gplx():
    doc = _new_doc()
    _header(doc, 'ĐƠN ĐỀ NGHỊ CẤP, ĐỔI GIẤY PHÉP LÁI XE',
            '(Mẫu 3, Phụ lục 29 kèm theo Thông tư số 12/2017/TT-BGTVT)')
    _add_para(doc, 'Kính gửi: Sở Giao thông Vận tải tỉnh/thành phố ..........................', sa=10)

    _add_para(doc, 'I. THÔNG TIN CÁ NHÂN', bold=True, sa=4)
    _add_field(doc, 'Họ và tên (chữ in hoa)', '{{ho_ten}}')
    _add_two(doc, 'Ngày, tháng, năm sinh', '{{ngay_sinh}}',
                  'Số CCCD/CMND', '{{cccd}}')
    _add_field(doc, 'Địa chỉ thường trú', '{{dia_chi_thuong_tru}}')
    _add_two(doc, 'Số điện thoại', '{{so_dien_thoai}}', 'Email', '{{email}}')

    _add_para(doc, '', sa=4)
    _add_para(doc, 'II. THÔNG TIN GIẤY PHÉP LÁI XE', bold=True, sa=4)
    _add_field(doc, 'Hạng GPLX đề nghị cấp/đổi', '{{hang_gplx_de_nghi}}')
    _add_two(doc, 'Số GPLX cũ (nếu đổi)', '{{so_gplx_cu}}',
                  'Ngày cấp', '{{ngay_cap_gplx_cu}}')
    _add_field(doc, 'Lý do đề nghị', '{{ly_do}}')

    _add_para(doc, '', sa=6)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(
        'Tôi cam đoan những nội dung khai trên đây là đúng sự thật. '
        'Nếu sai, tôi chịu hoàn toàn trách nhiệm trước pháp luật.'
    )
    _set_font(r, size=12)

    _add_para(doc, '', sa=6)
    _add_sign_row(doc, right_lines=[
        {'text': '{{dia_diem}}, ngày {{ngay}} tháng {{thang}} năm {{nam}}', 'italic': True},
        {'text': 'Người làm đơn', 'bold': True},
        {'text': '(Ký, ghi rõ họ tên)', 'italic': True},
        {'text': ''},{'text': ''},{'text': ''},
        {'text': '{{ho_ten}}', 'bold': True},
    ])
    _save(doc, 'mau-don-de-nghi-cap-doi-gplx.docx')


# ── 6. Đơn miễn giảm học phí ─────────────────────────────────────────────────

def build_mien_giam_hoc_phi():
    doc = _new_doc()
    _header(doc, 'ĐƠN ĐỀ NGHỊ MIỄN, GIẢM HỌC PHÍ',
            '(Kèm theo Nghị định số 81/2021/NĐ-CP ngày 27/8/2021 của Chính phủ)')
    _add_para(doc, 'Kính gửi: Ban Giám hiệu Trường .............................................', sa=10)

    _add_para(doc, 'I. THÔNG TIN HỌC SINH/SINH VIÊN', bold=True, sa=4)
    _add_field(doc, 'Họ và tên học sinh/sinh viên', '{{ho_ten_hoc_sinh}}')
    _add_two(doc, 'Ngày, tháng, năm sinh', '{{ngay_sinh}}',
                  'Lớp/Khóa học', '{{lop_khoa_hoc}}')
    _add_field(doc, 'Tên trường', '{{ten_truong}}')
    _add_field(doc, 'Địa chỉ thường trú', '{{dia_chi}}')
    _add_two(doc, 'Số điện thoại', '{{so_dien_thoai}}', 'Năm học', '{{nam_hoc}}')

    _add_para(doc, '', sa=4)
    _add_para(doc, 'II. THÔNG TIN PHỤ HUYNH', bold=True, sa=4)
    _add_field(doc, 'Họ và tên cha', '{{ho_ten_cha}}')
    _add_field(doc, 'Họ và tên mẹ', '{{ho_ten_me}}')

    _add_para(doc, '', sa=4)
    _add_para(doc, 'III. NỘI DUNG ĐỀ NGHỊ', bold=True, sa=4)
    _add_field(doc, 'Đối tượng miễn/giảm học phí', '{{doi_tuong}}')
    _add_field(doc, 'Lý do đề nghị miễn/giảm', '{{ly_do}}')

    _add_para(doc, '', sa=6)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(
        'Tôi xin cam kết những nội dung khai trên đây là đúng sự thật. '
        'Kính đề nghị nhà trường xem xét, giải quyết.'
    )
    _set_font(r, size=12)

    _add_para(doc, '', sa=6)
    _add_sign_row(doc, right_lines=[
        {'text': '{{dia_diem}}, ngày {{ngay}} tháng {{thang}} năm {{nam}}', 'italic': True},
        {'text': 'Người làm đơn', 'bold': True},
        {'text': '(Ký, ghi rõ họ tên)', 'italic': True},
        {'text': ''},{'text': ''},{'text': ''},
        {'text': '{{ho_ten_hoc_sinh}}', 'bold': True},
    ])
    _save(doc, 'mau-don-mien-giam-hoc-phi.docx')


# ── 7. Đơn hưởng BHXH một lần ────────────────────────────────────────────────

def build_bhxh_mot_lan():
    doc = _new_doc()
    _header(doc, 'ĐƠN ĐỀ NGHỊ HƯỞNG BẢO HIỂM XÃ HỘI MỘT LẦN',
            '(Mẫu 14-HSB kèm theo Quyết định số 222/QĐ-BHXH ngày 25/02/2021)')
    _add_para(doc, 'Kính gửi: Bảo hiểm xã hội ..................................................', sa=10)

    _add_para(doc, 'I. THÔNG TIN CÁ NHÂN', bold=True, sa=4)
    _add_field(doc, 'Họ và tên (chữ in hoa)', '{{ho_ten}}')
    _add_two(doc, 'Ngày, tháng, năm sinh', '{{ngay_sinh}}', 'Giới tính', '{{gioi_tinh}}')
    _add_two(doc, 'Số CCCD/CMND', '{{cccd}}', 'Số sổ BHXH', '{{so_so_bhxh}}')
    _add_field(doc, 'Địa chỉ thường trú', '{{dia_chi_thuong_tru}}')
    _add_field(doc, 'Số điện thoại', '{{so_dien_thoai}}')

    _add_para(doc, '', sa=4)
    _add_para(doc, 'II. THÔNG TIN BẢO HIỂM', bold=True, sa=4)
    _add_field(doc, 'Tên đơn vị cuối cùng tham gia BHXH', '{{don_vi_cuoi}}')
    _add_field(doc, 'Thời gian đóng BHXH (tháng)', '{{thoi_gian_dong}}')
    _add_field(doc, 'Lý do đề nghị hưởng một lần', '{{ly_do}}')

    _add_para(doc, '', sa=4)
    _add_para(doc, 'III. THÔNG TIN TÀI KHOẢN NHẬN TIỀN', bold=True, sa=4)
    _add_field(doc, 'Số tài khoản ngân hàng', '{{so_tai_khoan}}')
    _add_field(doc, 'Tên ngân hàng/Chi nhánh', '{{ten_ngan_hang}}')

    _add_para(doc, '', sa=6)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(
        'Tôi cam đoan những nội dung khai trên đây là đúng sự thật và chịu '
        'hoàn toàn trách nhiệm trước pháp luật.'
    )
    _set_font(r, size=12)

    _add_para(doc, '', sa=6)
    _add_sign_row(doc, right_lines=[
        {'text': '{{dia_diem}}, ngày {{ngay}} tháng {{thang}} năm {{nam}}', 'italic': True},
        {'text': 'Người làm đơn', 'bold': True},
        {'text': '(Ký, ghi rõ họ tên)', 'italic': True},
        {'text': ''},{'text': ''},{'text': ''},
        {'text': '{{ho_ten}}', 'bold': True},
    ])
    _save(doc, 'mau-don-huong-bhxh-mot-lan.docx')


# ── 8. Đơn đề nghị công nhận văn bằng ────────────────────────────────────────

def build_cong_nhan_van_bang():
    doc = _new_doc()
    _header(doc, 'ĐƠN ĐỀ NGHỊ CÔNG NHẬN VĂN BẰNG',
            '(Do cơ sở giáo dục nước ngoài cấp - Thông tư số 26/2021/TT-BGDĐT)')
    _add_para(doc, 'Kính gửi: .........................................................................', sa=10)

    _add_para(doc, 'I. THÔNG TIN CÁ NHÂN', bold=True, sa=4)
    _add_field(doc, 'Họ và tên (chữ in hoa)', '{{ho_ten}}')
    _add_two(doc, 'Ngày, tháng, năm sinh', '{{ngay_sinh}}', 'Giới tính', '{{gioi_tinh}}')
    _add_two(doc, 'Quốc tịch', '{{quoc_tich}}', 'Số CCCD/Hộ chiếu', '{{cccd_ho_chieu}}')
    _add_field(doc, 'Địa chỉ liên lạc', '{{dia_chi_lien_lac}}')
    _add_two(doc, 'Số điện thoại', '{{so_dien_thoai}}', 'Email', '{{email}}')

    _add_para(doc, '', sa=4)
    _add_para(doc, 'II. THÔNG TIN VĂN BẰNG', bold=True, sa=4)
    _add_field(doc, 'Tên văn bằng đề nghị công nhận', '{{ten_van_bang}}')
    _add_field(doc, 'Ngành/Chuyên ngành', '{{nganh_chuyen_nganh}}')
    _add_field(doc, 'Tên trường cấp văn bằng', '{{ten_truong}}')
    _add_two(doc, 'Nước cấp văn bằng', '{{nuoc_cap}}', 'Năm tốt nghiệp', '{{nam_tot_nghiep}}')

    _add_para(doc, '', sa=6)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(
        'Tôi cam đoan những thông tin kê khai trên đây là đúng sự thật. '
        'Kính đề nghị Quý cơ quan xem xét, giải quyết.'
    )
    _set_font(r, size=12)

    _add_para(doc, '', sa=6)
    _add_sign_row(doc, right_lines=[
        {'text': '{{dia_diem}}, ngày {{ngay}} tháng {{thang}} năm {{nam}}', 'italic': True},
        {'text': 'Người làm đơn', 'bold': True},
        {'text': '(Ký, ghi rõ họ tên)', 'italic': True},
        {'text': ''},{'text': ''},{'text': ''},
        {'text': '{{ho_ten}}', 'bold': True},
    ])
    _save(doc, 'mau-don-de-nghi-cong-nhan-van-bang.docx')


# ── 9. Đơn xin thôi quốc tịch ────────────────────────────────────────────────

def build_thoi_quoc_tich():
    doc = _new_doc()
    _header(doc, 'ĐƠN XIN THÔI QUỐC TỊCH VIỆT NAM',
            '(Kèm theo Luật Quốc tịch Việt Nam số 24/2008/QH12 và Luật sửa đổi số 56/2014/QH13)')
    _add_para(doc, 'Kính gửi: Chủ tịch nước Cộng hoà xã hội chủ nghĩa Việt Nam', sa=2)
    _add_para(doc, '         (Qua Bộ Tư pháp / Cơ quan đại diện Việt Nam ở nước ngoài)', sa=10)

    _add_para(doc, 'I. THÔNG TIN CÁ NHÂN', bold=True, sa=4)
    _add_field(doc, 'Họ và tên (chữ in hoa)', '{{ho_ten}}')
    _add_field(doc, 'Họ và tên khai sinh (nếu khác)', '{{ho_ten_khai_sinh}}')
    _add_two(doc, 'Ngày, tháng, năm sinh', '{{ngay_sinh}}', 'Giới tính', '{{gioi_tinh}}')
    _add_two(doc, 'Số CCCD/Hộ chiếu', '{{cccd_ho_chieu}}',
                  'Quốc tịch hiện tại', '{{quoc_tich_hien_tai}}')
    _add_field(doc, 'Địa chỉ thường trú tại Việt Nam', '{{dia_chi_vn}}')
    _add_field(doc, 'Địa chỉ cư trú ở nước ngoài (nếu có)', '{{dia_chi_nuoc_ngoai}}')

    _add_para(doc, '', sa=4)
    _add_para(doc, 'II. NỘI DUNG ĐỀ NGHỊ', bold=True, sa=4)
    _add_field(doc, 'Quốc tịch xin nhập', '{{quoc_tich_xin_nhap}}')
    _add_field(doc, 'Lý do xin thôi quốc tịch Việt Nam', '{{ly_do}}')

    _add_para(doc, '', sa=6)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(
        'Tôi xin cam đoan những nội dung khai trên đây là đúng sự thật và hoàn toàn '
        'chịu trách nhiệm trước pháp luật về tính chính xác của các thông tin đã khai.'
    )
    _set_font(r, size=12)

    _add_para(doc, '', sa=6)
    _add_sign_row(doc, right_lines=[
        {'text': '{{dia_diem}}, ngày {{ngay}} tháng {{thang}} năm {{nam}}', 'italic': True},
        {'text': 'Người làm đơn', 'bold': True},
        {'text': '(Ký, ghi rõ họ tên)', 'italic': True},
        {'text': ''},{'text': ''},{'text': ''},
        {'text': '{{ho_ten}}', 'bold': True},
    ])
    _save(doc, 'mau-don-xin-thoi-quoc-tich.docx')


# ── Main ──────────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    os.chdir(os.path.join(os.path.dirname(__file__), '..'))
    print('Rebuilding bad templates...\n')
    build_ket_hon()
    build_khai_tu()
    build_thuong_tru()
    build_dang_ky_thue()
    build_gplx()
    build_mien_giam_hoc_phi()
    build_bhxh_mot_lan()
    build_cong_nhan_van_bang()
    build_thoi_quoc_tich()
    print('\nDone! 9 templates rebuilt.')
