"""
create_missing_templates.py
---------------------------
Tạo file .docx stub cho các template được map trong DB nhưng chưa có trên disk.

Chạy: python scripts/create_missing_templates.py
"""
import os
import sys

if sys.stdout.encoding and sys.stdout.encoding.lower() not in ('utf-8', 'utf8'):
    sys.stdout.reconfigure(encoding='utf-8')

BASE_DIR = os.path.join(os.path.dirname(__file__), '..')
TEMPLATES_DIR = os.path.join(BASE_DIR, 'data', 'templates')

MISSING_STUBS = {
    'mau-ban-khai-ca-nhan.docx': {
        'title': 'BẢN KHAI CÁ NHÂN',
        'subtitle': '(Mẫu số 01 kèm theo Quyết định số 3635/QĐ-BTP)',
        'fields': [
            ('Họ và tên',                       '{{ho_ten}}'),
            ('Ngày tháng năm sinh',             '{{ngay_sinh}}'),
            ('Giới tính',                       '{{gioi_tinh}}'),
            ('Dân tộc',                         '{{dan_toc}}'),
            ('Quốc tịch',                       '{{quoc_tich}}'),
            ('Nơi sinh',                        '{{noi_sinh}}'),
            ('Số CCCD/CMND',                    '{{cccd}}'),
            ('Ngày cấp',                        '{{ngay_cap_cccd}}'),
            ('Nơi cấp',                         '{{noi_cap_cccd}}'),
            ('Hộ khẩu thường trú',              '{{ho_khau_thuong_tru}}'),
            ('Địa chỉ hiện tại',                '{{dia_chi_hien_tai}}'),
            ('Nghề nghiệp',                     '{{nghe_nghiep}}'),
            ('Nơi làm việc',                    '{{noi_lam_viec}}'),
            ('Trình độ học vấn',                '{{trinh_do_hoc_van}}'),
            ('Tình trạng hôn nhân',             '{{tinh_trang_hon_nhan}}'),
            ('Ngày lập bản khai',               '{{ngay_lap_ban_khai}}'),
        ],
    },
    'mau-don-de-nghi-cap-doi-gplx.docx': {
        'title': 'ĐƠN ĐỀ NGHỊ CẤP, ĐỔI GIẤY PHÉP LÁI XE',
        'subtitle': '(Mẫu 3, Phụ lục 29 kèm theo Thông tư số 12/2017/TT-BGTVT)',
        'fields': [
            ('Họ và tên',                       '{{ho_ten}}'),
            ('Ngày tháng năm sinh',             '{{ngay_sinh}}'),
            ('Số CCCD/CMND',                    '{{cccd}}'),
            ('Địa chỉ thường trú',              '{{dia_chi_thuong_tru}}'),
            ('Số điện thoại',                   '{{so_dien_thoai}}'),
            ('Email',                           '{{email}}'),
            ('Hạng GPLX đề nghị',               '{{hang_gplx_de_nghi}}'),
            ('Số GPLX cũ (nếu đổi)',            '{{so_gplx_cu}}'),
            ('Ngày cấp GPLX cũ',               '{{ngay_cap_gplx_cu}}'),
            ('Lý do đề nghị',                   '{{ly_do}}'),
            ('Ngày làm đơn',                    '{{ngay_lam_don}}'),
        ],
    },
    'mau-don-de-nghi-cong-nhan-van-bang.docx': {
        'title': 'ĐƠN ĐỀ NGHỊ CÔNG NHẬN VĂN BẰNG DO CƠ SỞ GIÁO DỤC NƯỚC NGOÀI CẤP',
        'subtitle': '(Kèm theo Thông tư số 26/2021/TT-BGDĐT)',
        'fields': [
            ('Họ và tên',                       '{{ho_ten}}'),
            ('Ngày tháng năm sinh',             '{{ngay_sinh}}'),
            ('Giới tính',                       '{{gioi_tinh}}'),
            ('Quốc tịch',                       '{{quoc_tich}}'),
            ('Số CCCD/Hộ chiếu',                '{{cccd_ho_chieu}}'),
            ('Địa chỉ liên lạc',                '{{dia_chi_lien_lac}}'),
            ('Số điện thoại',                   '{{so_dien_thoai}}'),
            ('Email',                           '{{email}}'),
            ('Tên văn bằng đề nghị công nhận',  '{{ten_van_bang}}'),
            ('Ngành/Chuyên ngành',              '{{nganh_chuyen_nganh}}'),
            ('Tên trường cấp văn bằng',         '{{ten_truong}}'),
            ('Nước cấp văn bằng',               '{{nuoc_cap}}'),
            ('Năm tốt nghiệp',                  '{{nam_tot_nghiep}}'),
            ('Ngày làm đơn',                    '{{ngay_lam_don}}'),
        ],
    },
    'mau-don-huong-bhxh-mot-lan.docx': {
        'title': 'ĐƠN ĐỀ NGHỊ HƯỞNG BẢO HIỂM XÃ HỘI MỘT LẦN',
        'subtitle': '(Mẫu 14-HSB kèm theo Quyết định số 222/QĐ-BHXH)',
        'fields': [
            ('Họ và tên',                       '{{ho_ten}}'),
            ('Ngày tháng năm sinh',             '{{ngay_sinh}}'),
            ('Giới tính',                       '{{gioi_tinh}}'),
            ('Số CCCD/CMND',                    '{{cccd}}'),
            ('Số sổ BHXH',                      '{{so_so_bhxh}}'),
            ('Địa chỉ thường trú',              '{{dia_chi_thuong_tru}}'),
            ('Số điện thoại',                   '{{so_dien_thoai}}'),
            ('Tên đơn vị cuối cùng tham gia BHXH', '{{don_vi_cuoi}}'),
            ('Thời gian đóng BHXH (tháng)',     '{{thoi_gian_dong}}'),
            ('Số tài khoản ngân hàng',          '{{so_tai_khoan}}'),
            ('Tên ngân hàng',                   '{{ten_ngan_hang}}'),
            ('Lý do đề nghị hưởng một lần',     '{{ly_do}}'),
            ('Ngày làm đơn',                    '{{ngay_lam_don}}'),
        ],
    },
    'mau-don-xin-thoi-quoc-tich.docx': {
        'title': 'ĐƠN XIN THÔI QUỐC TỊCH VIỆT NAM',
        'subtitle': '(Kèm theo Luật Quốc tịch Việt Nam số 24/2008/QH12)',
        'fields': [
            ('Họ và tên',                       '{{ho_ten}}'),
            ('Họ và tên khai sinh',             '{{ho_ten_khai_sinh}}'),
            ('Ngày tháng năm sinh',             '{{ngay_sinh}}'),
            ('Giới tính',                       '{{gioi_tinh}}'),
            ('Số CCCD/Hộ chiếu',                '{{cccd_ho_chieu}}'),
            ('Quốc tịch hiện tại',              '{{quoc_tich_hien_tai}}'),
            ('Quốc tịch xin nhập',              '{{quoc_tich_xin_nhap}}'),
            ('Địa chỉ thường trú tại Việt Nam', '{{dia_chi_vn}}'),
            ('Địa chỉ cư trú ở nước ngoài',    '{{dia_chi_nuoc_ngoai}}'),
            ('Lý do xin thôi quốc tịch',        '{{ly_do}}'),
            ('Ngày làm đơn',                    '{{ngay_lam_don}}'),
        ],
    },
    'mau-to-khai-dang-ky-thue.docx': {
        'title': 'TỜ KHAI ĐĂNG KÝ THUẾ',
        'subtitle': '(Mẫu 05-ĐK-TCT kèm theo Thông tư số 105/2020/TT-BTC)',
        'fields': [
            ('Họ và tên',                       '{{ho_ten}}'),
            ('Ngày tháng năm sinh',             '{{ngay_sinh}}'),
            ('Giới tính',                       '{{gioi_tinh}}'),
            ('Quốc tịch',                       '{{quoc_tich}}'),
            ('Số CCCD/CMND/Hộ chiếu',           '{{cccd}}'),
            ('Địa chỉ thường trú',              '{{dia_chi_thuong_tru}}'),
            ('Địa chỉ hiện tại',                '{{dia_chi_hien_tai}}'),
            ('Số điện thoại',                   '{{so_dien_thoai}}'),
            ('Email',                           '{{email}}'),
            ('Loại thu nhập đăng ký',           '{{loai_thu_nhap}}'),
            ('Nguồn thu nhập chính',            '{{nguon_thu_nhap}}'),
            ('Mã số thuế cũ (nếu có)',          '{{ma_so_thue_cu}}'),
            ('Ngày khai',                       '{{ngay_khai}}'),
        ],
    },
    'mau-to-khai-thue-su-dung-dat.docx': {
        'title': 'TỜ KHAI THUẾ SỬ DỤNG ĐẤT PHI NÔNG NGHIỆP',
        'subtitle': '(Kèm theo Thông tư số 153/2011/TT-BTC)',
        'fields': [
            ('Tên người nộp thuế',              '{{ten_nguoi_nop_thue}}'),
            ('Mã số thuế',                      '{{ma_so_thue}}'),
            ('Số CCCD/CMND',                    '{{cccd}}'),
            ('Địa chỉ',                         '{{dia_chi}}'),
            ('Số điện thoại',                   '{{so_dien_thoai}}'),
            ('Địa chỉ thửa đất',                '{{dia_chi_thua_dat}}'),
            ('Tờ bản đồ số',                    '{{to_ban_do_so}}'),
            ('Thửa đất số',                     '{{thua_dat_so}}'),
            ('Diện tích (m²)',                  '{{dien_tich}}'),
            ('Mục đích sử dụng',                '{{muc_dich_su_dung}}'),
            ('Hạn mức đất (m²)',                '{{han_muc_dat}}'),
            ('Giá đất tính thuế (đ/m²)',        '{{gia_dat_tinh_thue}}'),
            ('Năm tính thuế',                   '{{nam_tinh_thue}}'),
            ('Ngày khai',                       '{{ngay_khai}}'),
        ],
    },
}


def create_stub_docx(form_info: dict, out_path: str):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(form_info['title'])
    run.bold = True
    run.font.size = Pt(14)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(form_info['subtitle'])
    sub_run.italic = True
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.columns[0].width = Cm(7)
    table.columns[1].width = Cm(10)

    hdr = table.rows[0].cells
    hdr[0].text = 'Thông tin'
    hdr[1].text = 'Giá trị (điền vào)'
    for cell in hdr:
        for para in cell.paragraphs:
            para.runs[0].bold = True

    for label, placeholder in form_info['fields']:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = placeholder

    doc.add_paragraph()

    sig_para = doc.add_paragraph()
    sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig_para.add_run('Người khai\n(Ký, ghi rõ họ tên)')

    doc.save(out_path)
    print(f'  ✓ Đã tạo: {os.path.basename(out_path)}')


def main():
    if not os.path.isdir(TEMPLATES_DIR):
        os.makedirs(TEMPLATES_DIR)

    print('=== Tạo các file template bị thiếu ===\n')
    created = 0
    for fname, form_info in MISSING_STUBS.items():
        out_path = os.path.join(TEMPLATES_DIR, fname)
        if os.path.exists(out_path):
            print(f'  BỎ QUA (đã tồn tại): {fname}')
            continue
        try:
            create_stub_docx(form_info, out_path)
            created += 1
        except ImportError:
            print('  ERROR: pip install python-docx')
            return
        except Exception as e:
            print(f'  ERROR {fname}: {e}')

    print(f'\n✓ Đã tạo {created} file template mới.')


if __name__ == '__main__':
    main()
