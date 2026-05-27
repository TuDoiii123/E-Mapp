"""
fix_broken_templates.py
-----------------------
Kiểm tra các file template bị hỏng (HTML disguised as .docx) và
tạo file .docx thay thế với nội dung tối thiểu dựa trên loại biểu mẫu.

Chạy: python scripts/fix_broken_templates.py [--check-only]
"""
import os
import sys

if sys.stdout.encoding and sys.stdout.encoding.lower() not in ('utf-8', 'utf8'):
    sys.stdout.reconfigure(encoding='utf-8')

import argparse

BASE_DIR = os.path.join(os.path.dirname(__file__), '..')
TEMPLATES_DIR = os.path.join(BASE_DIR, 'data', 'templates')


def is_html(path: str) -> bool:
    with open(path, 'rb') as f:
        magic = f.read(16)
    html_sigs = [b'<!DOCTYPE', b'<!doctype', b'<html', b'<HTML', b' \n<!DOC', b'\n<!DOC']
    return any(magic.lstrip()[:9].upper().startswith(sig.upper().lstrip()[:9]) for sig in html_sigs)


# ── Nội dung các biểu mẫu thay thế ────────────────────────────────────────────
STUB_FORMS = {
    'mau-to-khai-dang-ky-ket-hon.docx': {
        'title': 'TỜ KHAI ĐĂNG KÝ KẾT HÔN',
        'subtitle': '(Ban hành kèm theo Thông tư số 04/2020/TT-BTP)',
        'fields': [
            ('Họ và tên người thứ nhất',       '{{ho_ten_nguoi_thu_nhat}}'),
            ('Ngày tháng năm sinh',             '{{ngay_sinh_nguoi_1}}'),
            ('Dân tộc',                         '{{dan_toc_nguoi_1}}'),
            ('Quốc tịch',                       '{{quoc_tich_nguoi_1}}'),
            ('Số CCCD/CMND',                    '{{cccd_nguoi_1}}'),
            ('Nơi thường trú/tạm trú',          '{{dia_chi_nguoi_1}}'),
            ('Họ và tên người thứ hai',         '{{ho_ten_nguoi_thu_hai}}'),
            ('Ngày tháng năm sinh',             '{{ngay_sinh_nguoi_2}}'),
            ('Dân tộc',                         '{{dan_toc_nguoi_2}}'),
            ('Quốc tịch',                       '{{quoc_tich_nguoi_2}}'),
            ('Số CCCD/CMND',                    '{{cccd_nguoi_2}}'),
            ('Nơi thường trú/tạm trú',          '{{dia_chi_nguoi_2}}'),
            ('Ngày đăng ký',                    '{{ngay_dang_ky}}'),
            ('Nơi đăng ký kết hôn (UBND)',      '{{noi_dang_ky}}'),
        ],
    },
    'mau-to-khai-dang-ky-khai-tu.docx': {
        'title': 'TỜ KHAI ĐĂNG KÝ KHAI TỬ',
        'subtitle': '(Ban hành kèm theo Thông tư số 04/2020/TT-BTP)',
        'fields': [
            ('Họ và tên người khai',            '{{ho_ten_nguoi_khai}}'),
            ('Số CCCD/CMND người khai',         '{{cccd_nguoi_khai}}'),
            ('Quan hệ với người chết',          '{{quan_he}}'),
            ('Họ và tên người đã chết',         '{{ho_ten_nguoi_chet}}'),
            ('Ngày tháng năm sinh',             '{{ngay_sinh_nguoi_chet}}'),
            ('Giới tính',                       '{{gioi_tinh}}'),
            ('Dân tộc',                         '{{dan_toc}}'),
            ('Quốc tịch',                       '{{quoc_tich}}'),
            ('Nơi thường trú/tạm trú',          '{{dia_chi_thuong_tru}}'),
            ('Ngày chết (ghi theo giấy chứng tử)', '{{ngay_chet}}'),
            ('Nơi chết',                        '{{noi_chet}}'),
            ('Nguyên nhân chết',                '{{nguyen_nhan_chet}}'),
            ('Ngày đăng ký khai tử',            '{{ngay_dang_ky}}'),
        ],
    },
    'mau-to-khai-dang-ky-thuong-tru.docx': {
        'title': 'TỜ KHAI ĐĂNG KÝ THƯỜNG TRÚ',
        'subtitle': '(Mẫu HK01 - Ban hành kèm theo Thông tư số 56/2021/TT-BCA)',
        'fields': [
            ('Họ và tên',                       '{{ho_ten}}'),
            ('Ngày tháng năm sinh',             '{{ngay_sinh}}'),
            ('Giới tính',                       '{{gioi_tinh}}'),
            ('Dân tộc',                         '{{dan_toc}}'),
            ('Số CCCD/CMND',                    '{{cccd}}'),
            ('Địa chỉ nơi ở hiện tại',          '{{dia_chi_hien_tai}}'),
            ('Địa chỉ đăng ký thường trú mới',  '{{dia_chi_thuong_tru_moi}}'),
            ('Chủ hộ (tên, quan hệ)',           '{{chu_ho_va_quan_he}}'),
            ('Lý do thay đổi',                  '{{ly_do_thay_doi}}'),
            ('Ngày đề nghị',                    '{{ngay_de_nghi}}'),
        ],
    },
    'mau-don-mien-giam-hoc-phi.docx': {
        'title': 'ĐƠN ĐỀ NGHỊ MIỄN, GIẢM HỌC PHÍ',
        'subtitle': '(Kèm theo Nghị định số 81/2021/NĐ-CP)',
        'fields': [
            ('Họ và tên học sinh/sinh viên',    '{{ho_ten_hoc_sinh}}'),
            ('Ngày tháng năm sinh',             '{{ngay_sinh}}'),
            ('Lớp/Khóa học',                    '{{lop_khoa_hoc}}'),
            ('Tên trường',                      '{{ten_truong}}'),
            ('Họ và tên cha',                   '{{ho_ten_cha}}'),
            ('Họ và tên mẹ',                    '{{ho_ten_me}}'),
            ('Địa chỉ thường trú',              '{{dia_chi}}'),
            ('Số điện thoại liên hệ',           '{{so_dien_thoai}}'),
            ('Đối tượng miễn/giảm',             '{{doi_tuong}}'),
            ('Lý do đề nghị miễn/giảm',        '{{ly_do}}'),
            ('Năm học',                         '{{nam_hoc}}'),
            ('Ngày làm đơn',                    '{{ngay_lam_don}}'),
        ],
    },
}


def create_stub_docx(form_info: dict, out_path: str):
    """Tạo file .docx tối thiểu với tiêu đề và bảng trường thông tin."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL

    doc = Document()

    # Căn lề trang
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2)

    # Tiêu đề
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(form_info['title'])
    run.bold = True
    run.font.size = Pt(14)

    # Phụ đề
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(form_info['subtitle'])
    sub_run.italic = True
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()  # khoảng trống

    # Bảng trường thông tin
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.columns[0].width = Cm(7)
    table.columns[1].width = Cm(10)

    # Header row
    hdr = table.rows[0].cells
    hdr[0].text = 'Thông tin'
    hdr[1].text = 'Giá trị (điền vào)'
    for cell in hdr:
        for para in cell.paragraphs:
            para.runs[0].bold = True

    for label, placeholder in form_info['fields']:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = placeholder

    doc.add_paragraph()

    # Chữ ký
    sig_para = doc.add_paragraph()
    sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig_para.add_run('Người khai\n(Ký, ghi rõ họ tên)')

    doc.save(out_path)
    print(f'  ✓ Đã tạo: {os.path.basename(out_path)}')


def main():
    parser = argparse.ArgumentParser(description='Fix broken HTML-disguised .docx templates')
    parser.add_argument('--check-only', action='store_true', help='Chỉ kiểm tra, không ghi file')
    args = parser.parse_args()

    if not os.path.isdir(TEMPLATES_DIR):
        print(f'Không tìm thấy thư mục templates: {TEMPLATES_DIR}')
        return

    print('=== Kiểm tra các file template bị hỏng ===\n')

    html_files = []
    for fname in sorted(os.listdir(TEMPLATES_DIR)):
        if not fname.endswith(('.doc', '.docx')):
            continue
        fpath = os.path.join(TEMPLATES_DIR, fname)
        try:
            if is_html(fpath):
                html_files.append(fname)
                status = '[CÓ STUB]' if fname in STUB_FORMS else '[KHÔNG CÓ STUB]'
                print(f'  HTML disguised: {fname}  {status}')
        except Exception as e:
            print(f'  ERROR reading {fname}: {e}')

    if not html_files:
        print('  Không có file bị hỏng nào.')
        return

    print(f'\nTổng: {len(html_files)} file HTML, {sum(1 for f in html_files if f in STUB_FORMS)} có stub định nghĩa sẵn')

    if args.check_only:
        print('\n[CHECK ONLY] Không ghi file.')
        return

    print('\n=== Tạo file thay thế ===\n')

    created = 0
    for fname, form_info in STUB_FORMS.items():
        fpath = os.path.join(TEMPLATES_DIR, fname)
        if not os.path.exists(fpath) or is_html(fpath):
            try:
                create_stub_docx(form_info, fpath)
                created += 1
            except ImportError:
                print('  ERROR: python-docx chưa được cài (pip install python-docx)')
                return
            except Exception as e:
                print(f'  ERROR tạo {fname}: {e}')
        else:
            print(f'  OK (bỏ qua): {fname}')

    print(f'\n✓ Đã tạo/thay thế {created} file template.')


if __name__ == '__main__':
    main()
